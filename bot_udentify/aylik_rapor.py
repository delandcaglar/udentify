from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx2pdf import convert
import os
import io
#import image_corpingv1
import heatmap
from datetime import date
import calendar
import datetime
import time
import calendar
#import docx_svg

import request1
from scipy.stats.stats import pearsonr





global_test =  False






#Genel_variabler
"""
Tüm_hafta_uzunlugu = son tarih - ilk tarih

"""
##pdf e cevirirken access vermen gerekli pythona dosyalar icin
firma = "Under Armour"
magza_statik_dosya_location_ismi = "Under Armour Akasya"
gereken_dosya_ismi = "Under Armour Akasya"
magaza_id_no = 240
magaza_adi = "Kadikoy"

ilk_tarih = "01/12/2020"
son_tarih = "25/12/2020"
tarih = ""

if global_test == True:
    print("global version")
    BASE_DIR = os.path.dirname ( os.path.dirname ( os.path.abspath ( __file__ ) ) )
    print ( BASE_DIR )
    magza_statik_dosya_location = os.path.join ( BASE_DIR, f"bot_udentify/firms/{firma}/Statik {magza_statik_dosya_location_ismi}" )
    print ( magza_statik_dosya_location )
    dosya_yolu = os.path.join ( BASE_DIR, f"bot_udentify/firms/{firma}" )
    magaza_dosyasi_ismi = f"{dosya_yolu}/{magza_statik_dosya_location_ismi}.docx"  ##bunu sona baglayacan
else:

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    print(BASE_DIR)
    magza_statik_dosya_location = os.path.join(BASE_DIR, 'bot_udentify/demo_resimler')
    print(magza_statik_dosya_location)
    dosya_yolu = os.path.join ( BASE_DIR, f"bot_udentify" )
    magaza_dosyasi_ismi = (f"{dosya_yolu}/demo.docx")



##global liste fonksiyonlari

def cal_average(num):
    sum_num = 0
    for t in num:
        sum_num = sum_num + t

    avg = sum_num / len(num)
    return avg




#1.1 için fonksiyon______________________


#1.1 global variables



""" 1-)Hangi günün pazartesi olduğunu bul 
2-)pazartesiden  sonra haftalara böl 4 parçada bak 5 günün ortalamasını al sonra kalan 2 günün ortalamasını al
3-)kişi sayısı 2 tanesinde fazla isi artmaktadır de 4 taneesi uzerinden
4-)ortalama süreeyi çek veye listelerin ortalamasını al ve yazdır 
5-)yukarıda yaptığun hafta içi hafta sonu değerlerini oranla 
6-)Yogunlik ortalamasi al toplam sureye bol


Kullanılacak varıabler *****
Magza_adi =
Magza_tarihi =
baslik_adlari = [] ## bunu dinamik liste yap"""




###1.11   tarih gerekli

def tarihi_bul(ilk_tarih1, son_tarih1):
    global tarih
    global ilk_tarih
    global son_tarih
    tarih = f"{ilk_tarih1} - {son_tarih1}" ##gloabal variable unutma

def magaza_adi_(magaza_adi1):
    global magaza_adi
    magaza_adi = magaza_adi1

tarihi_bul(ilk_tarih, son_tarih)
magaza_adi_(magza_statik_dosya_location_ismi)









#2.1 için fonksiyon______________________
""" 1-)Hangi günün pazartesi olduğunu bul 
2-)pazartesiden  sonra haftalara böl 4 parçada bak 5 günün ortalamasını al sonra kalan 2 günün ortalamasını al
3-)kişi sayısı 2 tanesinde fazla isi artmaktadır de 4 taneesi uzerinden
4-)ortalama süreeyi çek veye listelerin ortalamasını al ve yazdır 
5-)yukarıda yaptığun hafta içi hafta sonu değerlerini oranla 
6-)Yogunlik ortalamasi al toplam sureye bol

Gereken Variablelar ****
ortama_hafta_icleri = [a,b,c,d]
ortalama_hafta_sonlari = [a1,b1,c1,d1]
ortalama_sure = [a+b+c+d+a1+b1+c1+d1/8]  ##haftaya ölçeklendirilmiş veya tüm stringleri böl
ziyaretci_sayi_orani = [a+b+c+d /  a1+b1+c1+d1]
yogunluk_ortalamasi = yogunluk tablosunu al hepsini bol
yogunluk_grafiği_turevi = turevi fazla olan tarihi yaz

Kullanılacak varıabler *****
Tüm_hafta_uzunlugu ="""






gunler = request1.main1_1_tarih(240, ilk_tarih , son_tarih)
gunler_ve_musteri = request1.main1_1_musteri_sayisi(240, ilk_tarih , son_tarih)
gunler_ve_sure = request1.main1_1_musteri_suresi(240, ilk_tarih , son_tarih)
hafta_sonu_degerleri = []
hafta_ici_degerleri = []
hafta_ici_ortalama = float()
hafta_sonu_ortalama = float()
gunluk_ortalama_sure =float()



def findDay(date):
    born = datetime.datetime.strptime(date, "%d/%m/%Y").weekday()
    return (calendar.day_name[born])


def gunleri_classifiye_et():
    global hafta_sonu_degerleri
    global hafta_ici_degerleri
    global gunler
    for gun in gunler:
        if (findDay(gun) == "Sunday") or (findDay(gun) == "Saturday"):
            #print(gun)
            index = gunler.index ( gun )  # sifirdan basliyor
            #print ( 'The index of e:', index )
            #print("bugun sunday")
            hafta_sonu_degerleri.append ( index )
        else:
            #print(gun)
            index = gunler.index ( gun )  # sifirdan basliyor
            #print ( 'The index of e:', index )
            hafta_ici_degerleri.append(index)

#print(hafta_ici_degerleri)
#print(hafta_sonu_degerleri)


def hafta_ici_hafta_sonu_karsilastirma():
    global hafta_sonu_ortalama
    global hafta_ici_ortalama
    sabit = float ( 0 )
    print("hafta ici ve hafta sonu ortalamasi______")
    hafta_ici_ortalama_listesi = []
    hafta_sonu_ortalama_listesi = []
    for hafta_ici_degerler in hafta_ici_degerleri:
         #print(gunler_ve_musteri[hafta_ici_degerler])
         hafta_ici_ortalama_listesi.append(gunler_ve_musteri[hafta_ici_degerler])

    hafta_ici_ortalama = cal_average(hafta_ici_ortalama_listesi)
    print(hafta_ici_ortalama)
    for hafta_sonu_degerler in hafta_sonu_degerleri:
        #print ( gunler_ve_musteri[hafta_sonu_degerler] )
        hafta_sonu_ortalama_listesi.append ( gunler_ve_musteri[hafta_sonu_degerler]  )
    hafta_sonu_ortalama = cal_average ( hafta_sonu_ortalama_listesi )
    print ( hafta_sonu_ortalama )

def hafta_ici_hafta_sonu_oran_karsilastirmasi():
    global hafta_sonu_ortalama
    global hafta_ici_ortalama
    return (hafta_sonu_ortalama / hafta_ici_ortalama)


def hafta_ici_hafta_sonu_buyukluk_karsilastirmasi():
    global hafta_sonu_ortalama
    global hafta_ici_ortalama
    if hafta_sonu_ortalama > hafta_ici_ortalama:
        return "artmaktadır."
    elif hafta_sonu_ortalama < hafta_ici_ortalama:
        return "azalmaktadır."
    elif hafta_sonu_ortalama > hafta_ici_ortalama:
        return "hafta içi ile ortalamada günlük olarak aynı kalmaktadır."


## bbunlarin hepsinni fonksiyon icine koy calistir global degerlerini koynayi unnutma




gunleri_classifiye_et()
hafta_ici_hafta_sonu_karsilastirma() #ilk calismasi gereken fonksiyon !!!!!!!!! bbu part icin
hafta_ici_hafta_sonu_musteri_orani = ("{:.2f}".format(round(hafta_ici_hafta_sonu_oran_karsilastirmasi(), 2)))
hafta_ici_hafta_sonu_buyukluk_karsilastirma = hafta_ici_hafta_sonu_buyukluk_karsilastirmasi()





















def gunluk_ortalama_sure_(): ##tum gunluk sureleri topla gunluk sure sayisina bol
    global gunluk_ortalama_sure #sundan bir emin ol
    toplam_deger = float(0)
    for gunler in gunler_ve_sure:
        toplam_deger = toplam_deger + gunler
    return toplam_deger/(len(gunler_ve_sure))

print(gunluk_ortalama_sure_())
gunluk_ortalama_sure = ("{:.2f}".format(round(gunluk_ortalama_sure_(), 2))) ##


## density

density_tarihler = []
density_ortalama = []
density_degerler = []







#2.2 için fonksiyon______________________
""" 1-) en yuksek deger ile saatlik_sure degerini eslestir
2-) saatleri sabah 10-12  oglen 13-17  aksam 17-22  degerleerine bak son deger ortalamanin uzerinde ise
artislar kismina ekle 

Gereken Variablelar ****
sabah_degerleri = [a,b,c,]
oglen_degerleri = [a1,b1,c1,d1]
aksam_degerleri = [a1,b1,c1,d1,e1]

ortalama_hafta_sonlari = [a1,b1,c1,d1]
ortalama_sure = [a+b+c+d+a1+b1+c1+d1/8]  ##haftaya ölçeklendirilmiş veya tüm stringleri böl
ziyaretci_sayi_orani = [a+b+c+d /  a1+b1+c1+d1]
yogunluk_ortalamasi = yogunluk tablosunu al hepsini bol
yogunluk_grafiği_turevi = turevi fazla olan tarihi yaz

Kullanılacak varıabler *****
Tüm_hafta_uzunlugu ="""



sabah_degerleri = True
oglen_degerleri = True
aksam_degerleri = True

musteri_artisi_degeri = f""
def musteri_artisi():
    global musteri_artisi_degeri
    if sabah_degerleri == True:
        if oglen_degerleri == True:
            if aksam_degerleri == True :
                musteri_artisi_degeri =("sabah, öğlen ve akşam")
            else:
                musteri_artisi_degeri = ("sabah ve öğlen")
        else:
            if aksam_degerleri == True :
                musteri_artisi_degeri =("sabah ve akşam")
            else:
                musteri_artisi_degeri = ("sabah")
    else:
        if oglen_degerleri == True:
            if aksam_degerleri == True :
                musteri_artisi_degeri =("öğlen ve akşam")
            else:
                musteri_artisi_degeri = ("öğlen")
        else:
            if aksam_degerleri == True :
                musteri_artisi_degeri =("akşam")
            else:
                musteri_artisi_degeri = (" error atris yoktur") ##kontrol ett





#2.3 için fonksiyon______________________

""" tum alanlari 3 li list halinde siniflandir isim, yon, yogunluk degeri, (2.5 icin gerekli)
1-) Tum alanlara bak en yogun 3 unu sec
2-) Tum alanlara bak en yogun 4 unu sec
3-) tum alanlarin yukarisi, asagisi , solu, sagi diye siniflandir grafigin 
4-) yonlerden el fazlasina sahip olanini sec

Gereken Variablelar ****
yukari_degerleri = [a,b,c,]
sol_degerleri = [a1,b1,c1,d1]
sag_degerleri = [a1,b1,c1,d1,e1]
asagi_degerleri = [a1,b1,c1,d1,e1]
en_yogun_3_lu_degerler = [a9,b9,c9]
en_az_yogun_4_lu_degerler = [a0,b0,c0,d0]"""


yogunluk_haritasi_listesi =  [["FTW","sol alt",5  ],["Recover","sol alt",9 ],["Golf","sag alt",7 ]]

isim_listesi = []
konum_listesi = []
density_listesi = []
density_listesi_sorted = []

for elements in yogunluk_haritasi_listesi:
    isim_listesi.append(elements[0])
    konum_listesi.append(elements[1])
    density_listesi.append ( elements[2] )

density_listesi_sorted = density_listesi

density_listesi_sorted.sort()
print(density_listesi)


print(isim_listesi)

#2.4 için fonksiyon______________________
""" Buna bir geerekmiyor"""


#2.5 için fonksiyon______________________
""" 1-) tum alanlari 3 li list halinde siniflandir isim, yon, yogunluk degeri, 
 
 1-) Tum alanlara bak yogunlugu 1 in altindakileri sirala
2-) Tum alanlara bak yogunlugu 1 in ustundekiler sirala
3-) tum alanlarin yukarisi, asagisi , solu, sagi diye siniflandir grafigin 
4-) yonlerden el fazlasina sahip olanini sec

Gereken Listeler ****

Gereken_list_ler = [ [elma,1.8, yukari],[elma2,1.9, asagi] ]

Gereken Variablelar ****
yogunlugu_1_in_altindakiler = []
yogunlugu_1_in_ustundekiler = []


not: alani veya alanlarini ayarla cumle sonunda duzgun dursun"""

yogunlugu_1_in_altindakiler = [1]
yogunlugu_1_in_ustundekiler = []





#2.6 için fonksiyon______________________

""" 1-) Buraya bir open CV Atilabilinir veya elle yazilabilinir 
2-) yukari yonununde, asagi yonunde, sol yonunde, sag yonunde


not: alani veya alanlarini ayarla cumle sonunda duzgun dursun"""

hm1 = heatmap.find_where_is_everyone(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1_deneme.png"))
print('hadi bakalim')
print(hm1)
hm2 = heatmap.find_where_is_everyone(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1_deneme.png"))
print('hadi bakalim')
print(hm2)
hm3 = heatmap.find_where_is_everyone(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1_deneme.png"))
print('hadi bakalim')
print(hm3)
hm4 = heatmap.find_where_is_everyone(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1_deneme.png"))
print('hadi bakalim')
print(hm4)
hm5 = heatmap.find_where_is_everyone(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1_deneme.png"))
print('hadi bakalim')
print(hm5)
hm6 = heatmap.find_where_is_everyone(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1_deneme.png"))
print('hadi bakalim')
print(hm6)



#2.7 için fonksiyon______________________
""" 1-) listeleri_al = 3 uncu degeri yani degisimi en fazla olani koy(degisim)
2-)listeleri_al = 7 uncu degeri yani saniye en fazla olani koy(gecirilen sure)
3-)listeleri_al = 7 uncu degeri yani 15s giren en fazla olani koy(15s giren)

boyle gotur zaten tek fonksiyon

Gereken Listeler ****

Gereken_list_ler = [ [elma,1.8, yukari],[elma2,1.9, asagi] ]

Gereken Variablelar ****
yogunlugu_1_in_altindakiler = []
yogunlugu_1_in_ustundekiler = []


Kullanılacak varıabler *****
listeleri_al =
tarih =

not: alani veya alanlarini ayarla cumle sonunda duzgun dursun"""



#2.8 için fonksiyon______________________
"""1-)Hangi günün pazartesi olduğunu bul 
2-)pazartesiden  sonra haftalara böl 4 parçada bak 5 günün ortalamasını al sonra kalan 2 günün ortalamasını al
3-)kişi sayısı 2 tanesinde fazla isi artmaktadır de 4 taneesi uzerinden
4-)ortalama süreeyi çek veye listelerin ortalamasını al ve yazdır 
5-)yukarıda yaptığun hafta içi hafta sonu değerlerini oranla 
6-)Yogunlik ortalamasi al toplam sureye bol


Kullanılacak varıabler *****
Magza_alani =
Magza_tarihi ="""



#2.9 için fonksiyon______________________




#3.1 için fonksiyon______________________



##degerlerı refreshlemek ıcın gereklı scrıpt

def refresh():
    print('refresh baslatildi')







def start_writing_on_docx(firma,magza_statik_dosya_location_ismi,magaza_id_no,ilk_tarih,son_tarih,magaza_id_yogunluk,karsilastirilacak_isim_listesi):
    global global_test

    ##starting function
    if global_test == True:
        print ( "global version" )
        BASE_DIR = os.path.dirname ( os.path.dirname ( os.path.abspath ( __file__ ) ) )
        print ( BASE_DIR )
        magza_statik_dosya_location = os.path.join ( BASE_DIR,f"bot_udentify/firms/{firma}/Statik {magza_statik_dosya_location_ismi}" )
        print ( magza_statik_dosya_location )
        dosya_yolu = os.path.join ( BASE_DIR, f"bot_udentify/firms/{firma}" )
        magaza_dosyasi_ismi = f"{dosya_yolu}/{magza_statik_dosya_location_ismi}.docx"  ##bunu sona baglayacan
    else:

        BASE_DIR = os.path.dirname ( os.path.dirname ( os.path.abspath ( __file__ ) ) )
        print ( BASE_DIR )
        magza_statik_dosya_location = os.path.join ( BASE_DIR, 'bot_udentify/demo_resimler' )
        print ( magza_statik_dosya_location )
        dosya_yolu = os.path.join ( BASE_DIR, f"bot_udentify" )
        magaza_dosyasi_ismi = (f"{dosya_yolu}/demo.docx")


    #global variabllari kullanmayi unutma
    global tarih
    global magaza_adi


    tarihi_bul(ilk_tarih,son_tarih)
    magaza_adi_(magza_statik_dosya_location_ismi)

    print(magaza_id_no) ##burayi api ya bagladigin fonksiyonlari calistirmak icin kullanacaksin
    print(ilk_tarih)  ##bunlari norma diger tarihler icin kullanacaksin
    print(son_tarih)

    ##########___________________1.1 fonksiyonnlar

    global hafta_sonu_degerleri
    global hafta_ici_degerleri
    global hafta_sonu_ortalama
    global hafta_ici_ortalama
    global gunler
    global gunler_ve_musteri
    global gunler_ve_sure
    global hafta_ici_hafta_sonu_musteri_orani
    global hafta_ici_hafta_sonu_buyukluk_karsilastirma
    global gunluk_ortalama_sure

    gunler = request1.main1_1_tarih ( magaza_id_no, ilk_tarih, son_tarih )
    print('gunlerr')
    print(gunler)
    gunler_ve_musteri = request1.main1_1_musteri_sayisi ( magaza_id_no, ilk_tarih, son_tarih )
    print ( gunler_ve_musteri )
    gunler_ve_sure = request1.main1_1_musteri_suresi ( magaza_id_no, ilk_tarih, son_tarih )
    print ( gunler_ve_sure )

    gunleri_classifiye_et ()
    hafta_ici_hafta_sonu_karsilastirma ()  # ilk calismasi gereken fonksiyon !!!!!!!!! bbu part icin
    hafta_ici_hafta_sonu_musteri_orani = ("{:.2f}".format ( round ( hafta_ici_hafta_sonu_oran_karsilastirmasi (), 2 ) ))
    hafta_ici_hafta_sonu_buyukluk_karsilastirma = hafta_ici_hafta_sonu_buyukluk_karsilastirmasi ()
    gunluk_ortalama_sure = ("{:.2f}".format ( round ( gunluk_ortalama_sure_ (), 2 ) ))  ##

    print("bak bakalim")
    print(hafta_ici_hafta_sonu_musteri_orani)

    print('hafta sonu ortalama')
    print(hafta_sonu_ortalama)

    ##########___________________2.1 fonksiyonnlar

    yougunluk_miktarlari = request1.main_1_1_yogunluk_listesi ( magaza_id_no, ilk_tarih, son_tarih  )
    print ( yougunluk_miktarlari )
    yogunluk_tarihleri = request1.main_1_1_yogunluk_listesi_tarihler ( magaza_id_no, ilk_tarih, son_tarih  )
    print ( yogunluk_tarihleri )
    # dictionary = dict ( zip ( liste, liste1 ) )
    # print(dictionary)
    array_yogunluk = list ( zip ( yogunluk_tarihleri, yougunluk_miktarlari ) )
    print ( array_yogunluk )
    print ( request1.cal_average_density_arrray ( array_yogunluk ) )
    print ( request1.bigger_than_average_arrray ( array_yogunluk ) )
    ortalamanin_ustundeki_yogunluklar = request1.bigger_than_average_arrray ( array_yogunluk )
    print("bak____")
    print(ortalamanin_ustundeki_yogunluklar)

    ##sayfa-2.2
    liste = request1.main_2_1_saatler ( magaza_id_no, ilk_tarih, son_tarih )
    print ( liste )
    liste1 = request1.main_2_1_kisi_sure_saatlik ( magaza_id_no, ilk_tarih, son_tarih )
    print ( liste1 )
    array = list ( zip ( liste, liste1 ) )
    print ( array )
    sabah_listesi = array[0:3]
    sabah_ortalmasi = (request1.cal_average_array ( sabah_listesi ))
    oglen_listesi = array[3:7]
    oglen_ortalmasi = (request1.cal_average_array ( oglen_listesi ))
    aksam_listesi = array[7:12]
    aksam_ortalmasi = (request1.cal_average_array ( aksam_listesi ))
    en_yogun_zaman = ( request1.sabah_ogle_aksam ( sabah_ortalmasi, oglen_ortalmasi, aksam_ortalmasi ) )
    max_saat_araligi = ( request1.max_number_array ( array ) )

    #sayfa-2.3
    liste_yogunluk_top_5 = request1.main_2_3_yogunluk_haritasi_top_5(magaza_id_yogunluk,ilk_tarih, son_tarih )
    print("lolall")
    print(liste_yogunluk_top_5)

    liste_yogunluk_bottom_5 =request1.main_2_3_yogunluk_haritasi_bottom_5 ( magaza_id_yogunluk, ilk_tarih, son_tarih )
    print ( liste_yogunluk_bottom_5 )

    # sayfa-2.5
    yogunlugu_1_in_ustundekiler = request1.main_2_5_yogunluk_haritasi_orani_top_5_orani ( magaza_id_yogunluk, ilk_tarih, son_tarih )
    print ( "top_5" )
    print ( yogunlugu_1_in_ustundekiler )
    yogunlugu_1_in_altindakiler = request1.main_2_5_yogunluk_haritasi_orani_bottom_5_orani ( magaza_id_yogunluk, ilk_tarih, son_tarih )
    print ( "bot_5" )
    print ( yogunlugu_1_in_altindakiler )

    # sayfa-2.7


    performans_tablosu_listesi = request1.main_2_7_performans_tablosu_alansal ( magaza_id_no, ilk_tarih, son_tarih, "Count" )
    print ( "performans_tablosu_listesi" )
    print ( performans_tablosu_listesi )
    magaza_bolumu_id = (performans_tablosu_listesi[0][1])

    ##________

    kisi_sayisi2_7 = request1.main_2_7_performans_tablosu_alansal ( magaza_id_no, ilk_tarih, son_tarih, "PrevCount" )
    kisi_sayisi2_7_isim = kisi_sayisi2_7[0][0].title ()
    kisi_sayisi2_7_deger = kisi_sayisi2_7[1][5]

    gecirilen_sure_2_7 = request1.main_2_7_performans_tablosu_alansal ( magaza_id_no, ilk_tarih, son_tarih,
                                                                        "PrevDwell" )
    gecirilen_sure_2_7_isim = gecirilen_sure_2_7[0][0].title ()
    gecirilen_sure_2_7_deger = gecirilen_sure_2_7[1][7]

    ilgi_oranı_2_7 = request1.main_2_7_performans_tablosu_alansal ( magaza_id_no, ilk_tarih, son_tarih,
                                                                    "_15s_Enterance" )
    ilgi_oranı_2_7_isim = ilgi_oranı_2_7[0][0].title ()
    ilgi_oranı_2_7_deger = ilgi_oranı_2_7[1][8]

    ilgi_oranı_n2_7 = request1.main_2_7_performans_tablosu_alansal ( magaza_id_no, ilk_tarih, son_tarih,
                                                                     "n_15s_Enterance" )
    ilgi_oranı_n2_7_isim = ilgi_oranı_n2_7[0][0].title ()
    ilgi_oranı_n2_7_deger = ilgi_oranı_n2_7[1][8]

    en_cok_zaman_2_7 = request1.main_2_7_performans_tablosu_alansal ( magaza_id_no, ilk_tarih, son_tarih, "Dwell" )
    en_cok_zaman_2_7_isim = en_cok_zaman_2_7[0][0].title ()
    en_cok_zaman_2_7_deger = en_cok_zaman_2_7[1][6]

    # sayfa-2.7







    gunler_2_8 = request1.main_2_8_saatler_gunluk(magaza_bolumu_id, ilk_tarih, son_tarih)

    gunler_ve_musteri_2_8 = request1.main_2_8_kisi_miktari(magaza_bolumu_id, ilk_tarih, son_tarih)

    gunler_ve_sure_2_8 = request1.main_2_8_kisi_sure(magaza_bolumu_id, ilk_tarih, son_tarih)

    hafta_ici_degerler_2_8 , haftasonu_degerleri_2_8 = request1.gunleri_classifiye_et(gunler_2_8)[0],request1.gunleri_classifiye_et(gunler_2_8)[1]

    print("ilk_deneme")
    print(hafta_ici_degerler_2_8)
    print(haftasonu_degerleri_2_8)

    hafta_ici_ortalama_2_8, haftasonu_ortalama_2_8 = request1.hafta_ici_hafta_sonu_ortalamalar(hafta_ici_degerler_2_8,haftasonu_degerleri_2_8,gunler_ve_musteri_2_8)[0],request1.hafta_ici_hafta_sonu_ortalamalar(hafta_ici_degerler_2_8,haftasonu_degerleri_2_8,gunler_ve_musteri_2_8)[1]

    print(hafta_ici_ortalama_2_8)
    print(haftasonu_ortalama_2_8)

    katsayi_2_8 = request1.hafta_ici_hafta_sonu_oran_karsilastirmasi(hafta_ici_ortalama_2_8,haftasonu_ortalama_2_8)
    print(katsayi_2_8)

    artis_2_8 = request1.hafta_ici_hafta_sonu_buyukluk_karsilastirmasi(hafta_ici_ortalama_2_8,haftasonu_ortalama_2_8)

    ortalama_sure_2_8 = request1.gunluk_ortalama_sure_2_8(gunler_ve_sure_2_8)

    #### saatlik kisi sure grafigi

    saatler_listes_2_8_ = (request1.saatler_2_8_ ( magaza_bolumu_id, ilk_tarih, son_tarih ))
    kisi_sure_saatlik_listes_2_8_ = (request1.toplam_kisi_sure_saatlik_2_8_ ( magaza_bolumu_id, ilk_tarih, son_tarih ))

    array = list ( zip ( saatler_listes_2_8_, kisi_sure_saatlik_listes_2_8_ ) )
    artisin_oldugu_sure_2_8 = request1.max_number_array ( array )
    print ( artisin_oldugu_sure_2_8 )

    kisi_sure_saatlik_liste_2_8_ = (request1.kisi_sure_saatlik_2_8_ ( magaza_bolumu_id, ilk_tarih, son_tarih ))

    saat_sure_2_8 = request1.saat_sure_saatlik_2_8_ ( magaza_bolumu_id, ilk_tarih, son_tarih )

    r_degeri = (pearsonr ( kisi_sure_saatlik_liste_2_8_, saat_sure_2_8 ))

    korelasyon_degeri_2_8= float("{:.2f}".format(round(r_degeri[0], 3)))

    print("korelasyon_degeero")
    print(korelasyon_degeri_2_8)
    korelasyon_turu = request1.korelasyon_orani(korelasyon_degeri_2_8)

    #2_8 yogunluk
    yogunluk_tarihleri_2_8 = (request1.yogunluk_listesi_tarihler_2_8_ ( magaza_bolumu_id, ilk_tarih, son_tarih ))
    yougunluk_miktarlari_2_8 = (request1.yogunluk_listesi_2_8_ ( magaza_bolumu_id, ilk_tarih, son_tarih ))
    array_yogunluk_2_8 = list ( zip ( yogunluk_tarihleri_2_8, yougunluk_miktarlari_2_8 ) )
    print ( array_yogunluk_2_8 )
    print ( request1.cal_average_density_arrray ( array_yogunluk_2_8 ) )
    print ( request1.bigger_than_average_arrray ( array_yogunluk_2_8 ) )
    ortalamanin_ustundeki_yogunluklar_2_8 = request1.bigger_than_average_arrray ( array_yogunluk_2_8 )
    print ( "bak____" )
    print ( ortalamanin_ustundeki_yogunluklar_2_8 )

    yogunlugu_fazla_olan_tarihler_2_8 = (request1.gunleri_classifiye_et ( ortalamanin_ustundeki_yogunluklar_2_8 ))



    hafta_ici_hafta_sonu_oranlari_2_8 = request1.hafta_ici_hafta_sonu_oranlari(yogunlugu_fazla_olan_tarihler_2_8[0],yogunlugu_fazla_olan_tarihler_2_8[1])

    print ( "deneme11111122" )
    print ( performans_tablosu_listesi )
    edinme_hunisi_yuzde = request1.main_2_8_edinme_hunisi(performans_tablosu_listesi[0][3],ilk_tarih, son_tarih) # burayı a veya b ye goree ayarlayabılıyoruz
    print("deneme1111111") # ! burayi degistirdin 4 u yukarida 3 yaprin
    print(performans_tablosu_listesi)
    print(performans_tablosu_listesi[0][3])











    document = Document()


    #Header ekleme kodu
    header = document.sections[0].header
    htable=header.add_table(1, 2, Inches(13))
    htab_cells=htable.rows[0].cells
    ht0=htab_cells[1].add_paragraph()
    kh=ht0.add_run()
    kh.add_picture('/Users/ilkedelandcaglar/Downloads/udentify/bot_udentify/udentıfy.png', width=Inches(1))
    #ht1=htab_cells[1].add_paragraph('put your header text here')
    ht0.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ##syafa ekleme kodu
    def create_element(name):
        return OxmlElement(name)


    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)


    def add_page_number(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        page_run = paragraph.add_run()
        t1 = create_element('w:t')
        create_attribute(t1, 'xml:space', 'preserve')
        t1.text = 'Sayfa '
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = create_element('w:t')
        create_attribute(t2, 'xml:space', 'preserve')
        t2.text = ' / '
        of_run._r.append(t2)

        fldChar3 = create_element('w:fldChar')
        create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = create_element('w:instrText')
        create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = create_element('w:fldChar')
        create_attribute(fldChar4, 'w:fldCharType', 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)

    add_page_number(document.sections[0].footer.paragraphs[0])


    #sayfe-giris


    p = document.add_paragraph(f"")
    p = document.add_paragraph(f"")
    p = document.add_paragraph(f"")
    p = document.add_paragraph(f"")
    p = document.add_paragraph(f"")
    p = document.add_paragraph(f"")



    ##docx_svg._SvgParser.monkey ()



    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/under_armour_logo.png"))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"")
    p.add_run(f"({magaza_adi})").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER



    #sayfa-icindekiler_____________________________

    document.add_page_break()

    document.add_heading("İçindekiler", level=0)

    p = document.add_paragraph(f"")
    p.add_run(f"1.0 Giriş").bold = True
    p = document.add_paragraph(f"       1.1 Raporun Amacı")

    p = document.add_paragraph(f"")
    p.add_run(f"2.0 Undentify Genel Raporu").bold = True
    p = document.add_paragraph(f"       2.1 Günlük Kişi Süre Grafiği")
    p = document.add_paragraph(f"       2.2 Saatlik Kişi Süre Grafiği")
    p = document.add_paragraph(f"       2.3 Yoğunluk Haritası")
    p = document.add_paragraph(f"       2.4 Mağaza İçi Yoğunluk Dağılımı")
    p = document.add_paragraph(f"       2.5 Metre Kare Başına Düşen Yoğunluk Haritası")
    p = document.add_paragraph(f"       2.6 Isı Haritası")
    p = document.add_paragraph(f"       2.7 Performans Kıyaslaması")
    p = document.add_paragraph(f"       2.8 {performans_tablosu_listesi[0][0].title()}") #.capitilize da olur
    p = document.add_paragraph(f"       2.9 Kategori Karşılaştırması")

    p = document.add_paragraph(f"")
    p.add_run(f"3.0 Sonuç").bold = True
    p = document.add_paragraph(f"       3.1 Analizin sonucu")



    #sayfa-1.1_____________________________________________________

    document.add_page_break()
    document.add_heading('1.0 Giriş', 0)

    document.add_heading('1.1 Raporun Amacı', level=1)

    baslik_adlari = f"günlük kişi süre grafiği, saatlik kişi süre grafiği, yoğunluk haritası, mağaza içi yoğunluk dağılımı, metre kare başına düşen yoğunluk haritası, ısı haritası ,performans kıyaslaması, {performans_tablosu_listesi[0][0].title()}, kategori karşılaştırması"

    p = document.add_paragraph(f"Bu raporun amacı {magaza_adi} mağazasının belirtilen {tarih} aralığındaki performans metrikleri ({baslik_adlari}) üzerinde yapılan analizler sonucu elde edilen bilgiyi sağlamaktır.")
    #magzalarinin



    #sayfa-2.1_____________________________________________________

    document.add_page_break()
    document.add_heading('2.0 Udentify Genel Raporu', 0)






    """p = document.add_paragraph(f"")
    p.add_run(f"Günlük Kişi Süre Grafiği").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

    document.add_heading('2.1 Günlük Kişi Süre Grafiği', level=1)


    ee = os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/giren_m_1.png")
    ee_o = os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/toplam_kisi_sure.png")

    #deneme = image_corpingv1.image_corp_edit_infile(ee,ee_o,1008,286) #kullanmamis oldugundan emin old !!!!!!!!!!!!!___________________




    document.add_picture(ee_o, width=Inches(6) , height=Inches(2)) #kesme fonksiyonlari cok iyi degil
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Toplam Kişi & Ortalama Geçirilen Süre Grafiği)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER





    document.add_paragraph(
        f"Kişi-süre grafiği incelendiğinde {tarih} tarihleri arasında mağazaya gelen ziyaretçi sayısında ve geçirdikleri süredeki değişiklikler görülmektedir.", style='List Bullet'
    )

    document.add_paragraph(
        f"Hafta sonları kişi sayısı {hafta_ici_hafta_sonu_buyukluk_karsilastirma}", style='List Bullet'
    )

    document.add_paragraph(
        f"Mağazaya gelen ziyaretçilerin mağaza içerisinde geçirdikleri ortalama süre {gunluk_ortalama_sure} saniyedir.", style='List Bullet'
    )

    document.add_paragraph(
        f"Hafta sonu mağazaya gelen ortalama ziyaretçi sayısı, hafta içinin {hafta_ici_hafta_sonu_musteri_orani} katıdır.", style='List Bullet'
    )



    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/yogunluk_yogunluk_ortalamasi.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Yoğunluk Trendi Tablosu)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        'Yoğunluk grafiği incelendiğinde, yoğunluğun ortalamanın üstüne çıktığı tarihler görülmektedir.', style='List Bullet'
    )

    document.add_paragraph(
        f'{ortalamanin_ustundeki_yogunluklar[0].title()} tarihi yoğunluğun en çok arttığı tarihlerden biridir. Yoğunluktaki artışın sebebi, kişi sayısındaki artıştır.', style='List Bullet'
    )

    document.add_paragraph(
        f'Yoğunluğun en çok artış gösterdiği diğer bir tarih {ortalamanin_ustundeki_yogunluklar[0].title()}’dir. O gün yoğunluğun artmasının sebebi, mağaza içerisinde geçirilen sürenin artış göstermesidir.', style='List Bullet'
    )




    #sayfa-2.2_____________________________________________________
    #document.add_page_break()

    """p = document.add_paragraph(f"")
    p.add_run(f"Saatlik Kişi Süre Grafiği").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

    document.add_heading('2.2 Saatlik Kişi Süre Grafiği', level=1)


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/saatlik_kisi_sure.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Toplam Kişi - Ortama Geçirilen Saatlik Süre Grafiği)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        f'Mağazaya en çok sayıda ziyaretçinin geldiği saat aralığı {max_saat_araligi} saatleri arasıdır', style='List Bullet'
    )

    document.add_paragraph(
        f'Mağaza içerisinde geçirilen {en_yogun_zaman} saatlerinde artmaktadır. ', style='List Bullet'
    )

    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/en_cok_ziyaret.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(En Çok Vakit Geçirilen Veya En Çok Ziyaret Edilenler Tablosu)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER



    #sayfa-2.3_____________________________________________________
    document.add_page_break()

    """p = document.add_paragraph(f"")
    p.add_run(f"Yoğunluk Haritası").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

    document.add_heading('2.3 Yoğunluk Haritası', level=1)


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/yogunluk_haritasi.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Yoğunluk Haritaları Ve Tabloları)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/top5.png"), width=Inches(2) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        f'Yukarıdaki yoğunluk haritası incelendiğinde, mağazanın en yoğun alanlarının {liste_yogunluk_top_5[0].title()}, {liste_yogunluk_top_5[1].title()} ve {liste_yogunluk_top_5[2].title()} alanları olduğu görülür.', style='List Bullet'
    )

    document.add_paragraph(
        f'Yoğunluğun en az olduğu alanlar ise; {liste_yogunluk_bottom_5[0].title()}, {liste_yogunluk_bottom_5[1].title()}, {liste_yogunluk_bottom_5[2].title()}, {liste_yogunluk_bottom_5[3].title()}, ve {liste_yogunluk_bottom_5[4].title()} alanlarıdır.', style='List Bullet'
    )

    document.add_paragraph(
        f'Girişin sol tarafında kalan alanlar sağa göre daha yoğundur.', style='List Bullet'
    )









    #sayfa-2.4_____________________________________________________

    document.add_page_break()


    """p = document.add_paragraph(f"")
    p.add_run(f"Mağaza İçi Yoğunluk Dağılımı").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

    document.add_heading('2.4 Mağaza İçi Yoğunluk Dağılımı', level=1)


    p = document.add_paragraph(f"")
    p.add_run('Mağaza içerisindeki alanlar, yoğunluğu en büyük olandan en küçük olana doğru sıralanmıştır.').italic = True


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/magaza_ici_dagilimi.png"), width=Inches(6) , height=Inches(7))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Yoğunluk Alan Tablosu)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER




    #sayfa-2.5_____________________________________________________
    document.add_page_break()


    """p = document.add_paragraph(f"")
    p.add_run(f"Metre Kare Başına Düşen Yoğunluk Haritası").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

    document.add_heading('2.5 Metre Kare Başına Düşen Yoğunluk Haritası', level=1)


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/yogunlu_m2.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Yoğunluk Haritası Ve Tablosu)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        f"Metre kare başına düşen yoğunluklar incelendiğinde, bazı alanlarda bu oranın 1’den küçük olduğu görülür. Bu, birim metre kare başına birim yoğunluktan daha az yoğunluk düştüğünü gösterir.  Birçok alanda metre kare başına düşen yoğunluk 1’in altında kalmıştır. Örneğin; {yogunlugu_1_in_ustundekiler[0].title()}, {yogunlugu_1_in_ustundekiler[1].title()}, {yogunlugu_1_in_ustundekiler[2].title()} alanları en düşük metre kare başına düşen yoğunluk oranına sahiptir.", style='List Bullet'
    )

    document.add_paragraph(
        'Bu alanlarda, metre karenin küçültülmesi düşünülebilir. Buna karar vermek için metre kare kademeli olarak küçültülüp bunun satışa etkisi incelenmelidir. Alan küçülürken satış azalmıyorsa alan küçültülmelidir.', style='List Bullet'
    )
    try:
        listeyy = f"{yogunlugu_1_in_altindakiler[0].title ()}, {yogunlugu_1_in_altindakiler[1].title ()}, {yogunlugu_1_in_altindakiler[2].title ()}"
    except:
        try:
            listeyy = f"{yogunlugu_1_in_altindakiler[0].title ()}, {yogunlugu_1_in_altindakiler[1].title ()}"
        except:
            try:
                listeyy = f"{yogunlugu_1_in_altindakiler[0].title ()}"
            except:
                listeyy = f"hepsi cok iyi"


    document.add_paragraph(
        f"Bazı alanlarda ise bu oran 1’in üstündedir. Yani, birim metre kare başına birim yoğunluktan daha fazlası düşmüştür. Örneğin; {listeyy} alanları.", style='List Bullet'
    )

    document.add_paragraph(
        'Aynı mantıkla, küçültülen alanlardan kalan payın bu alana verilerek bu alanın genişletilmesi düşünülebilir. ', style='List Bullet'
    )


    #sayfa-2.6_____________________________________________________
    document.add_page_break()

    """p = document.add_paragraph(f"")
    p.add_run(f"Isı Haritası").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""


    document.add_heading('2.6 Isı Haritası', level=1)

    p = document.add_paragraph(f"Kırmızı: Çok yoğun Sarı: Orta yoğun Yeşil: Az yoğun")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        'Isı haritaları incelendiğinde, mağaza içerisinde yoğunluğu daha az ve daha fazla olan alanlar görülmektedir. ', style='List Bullet'
    )


    document.add_paragraph(
        f"Kasa kamerasının {hm1} alanınnda, basketball kamerasının {hm1} alanınnda, men’s run kamerasının {hm2} alanınnda, heat gear kamerasının {hm1} alanınnda, FTW kamerasının {hm1} alanınnda, giriş kamerasının {hm1} alanınnda yoğunluğun arttığı görülmektedir.", style='List Bullet'
    )

    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi1.png"), width=Inches(6) , height=Inches(1.9))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi2.png"), width=Inches(6) , height=Inches(1.9))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/isi3.png"), width=Inches(6) , height=Inches(1.9))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Isı Haritaları)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER




    #sayfa-2.7_____________________________________________________
    #document.add_page_break()


    """p = document.add_paragraph(f"")
    p.add_run(f"Performans Kıyaslaması").bold = True"""


    document.add_heading('2.7 Performans Kıyaslaması', level=1)


    p = document.add_paragraph(f"")
    p.add_run('*15/01/2020-28/01/2020 tarih aralığı, 01/01/2020-14/01/2020 aralığı ile kıyaslanmıştır.').italic = True


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/performans_tablosu.png"), width=Inches(6) , height=Inches(5))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"(Performans Tablosu)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER







    p = document.add_paragraph(f"")
    p.add_run('*Performans tablosu, mağaza içerisindeki yoğunluk payı en fazla olan bölgeden en az olan bölgeye doğru sıralanmıştır.').italic = True


    p = document.add_paragraph(f"Performans tablosu incelendiğinde;")

    document.add_paragraph(
        f'Kişi sayısı en çok artan alan %{kisi_sayisi2_7_deger} artış ile {kisi_sayisi2_7_isim} alanı,', style='List Bullet'
    )

    document.add_paragraph(
        f'Geçirilen süresi en çok artan alan %{gecirilen_sure_2_7_deger} artış ile {gecirilen_sure_2_7_isim} alanı,', style='List Bullet'
    )
    document.add_paragraph(
        f'İlgi oranı en yüksek olan alan %{ilgi_oranı_2_7_deger} ile {ilgi_oranı_2_7_isim} alanı,', style='List Bullet'
    )

    document.add_paragraph(
        f'İlgi oranı en az olan alan %{ilgi_oranı_n2_7_deger} ile {ilgi_oranı_n2_7_isim} alanı,', style='List Bullet'
    )
    document.add_paragraph(
        f'En çok zaman geçirilen alan ortalama {en_cok_zaman_2_7_deger} saniye ile {en_cok_zaman_2_7_isim} alanı,', style='List Bullet'
    )
    document.add_paragraph(
        f'En çok ziyaret edilen alan günde ortalama {performans_tablosu_listesi[1][3]} kişi ile {performans_tablosu_listesi[0][0].title()} alanı olmuştur.', style='List Bullet'
    )






    #sayfa-2.8_____________________________________________________
    #document.add_page_break()


    document.add_heading(f'2.8 {performans_tablosu_listesi[0][0].title()}', level=1)


    p = document.add_paragraph(f"Mağaza içerisindeki en yoğun alan olan {performans_tablosu_listesi[0][0].title()} alanının detayına inildiğinde;")

    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/kisi_sure-1.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"({performans_tablosu_listesi[0][0].title()} Alanı Toplam Kişi Ve Günlük Süre Grafiği)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        f'Günlük kişi süre grafiği incelendiğinde, hafta sonu mağazaya gelen ortalama kişi sayısı hafta sonuna göre {artis_2_8} Bu oran {katsayi_2_8} katıdır.', style='List Bullet'
    )

    document.add_paragraph(
        f'Bu alana gelen bir kişinin alanda geçirdiği ortalama süre {ortalama_sure_2_8} saniyedir.', style='List Bullet'
    )


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/kisi_sure_grafigi.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"({performans_tablosu_listesi[0][0].title()} Alanı Toplam Kişi Ve Saatlik Süre Grafiği)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        f'Saatlik kişi süre grafiği incelendiğinde mağazaya en çok sayıda ziyaretçinin geldiği zaman aralığının {artisin_oldugu_sure_2_8} saatleri arası olduğu görülür.', style='List Bullet'
    )

    document.add_paragraph (
        'Geçirilen sürenin arttığı saatlerde müşteri ilgisi artmaktadır.', style='List Bullet'
    )

    document.add_paragraph(
        f'Kişi sayısı ile süre korelasyonu {korelasyon_degeri_2_8} dir, yani {korelasyon_turu} vardır.', style='List Bullet'
    )


    document.add_paragraph(
        'Bu saatlerde, artan yoğunluk potansiyelini satışa dönüştürmek için personel ilgisine ekstra dikkat etmek gerekir.', style='List Bullet'
    )


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/yogunluk.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"({performans_tablosu_listesi[0][0].title()} Alanı Yoğunluk Ve Yoğunluk Ortalaması Grafiği)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        f'Yoğunluk grafiğinde, yoğunluğun ortalamanın üstünde ve altında kaldığı günler görülmektedir. Yoğuluğun arttığı tarihler {request1.list_to_string(ortalamanin_ustundeki_yogunluklar_2_8)} olarak gözükmektedir. Yoğunluğun arttığı günlerde, satışın da artması beklenir.', style='List Bullet'
    )

    document.add_paragraph(
        f'{hafta_ici_hafta_sonu_oranlari_2_8}', style='List Bullet'
    )

    ##-_____________________________________________________


    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/korelasyonlar.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"({performans_tablosu_listesi[0][0].title()} Alanı Korelasyon Tablosu)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    document.add_paragraph(
        'Korelasyonlar incelendiğinde, bu alana gelen ziyaretçilerin büyük çoğunluğunun bu alandan sonra men’s sportstyle alanına yöneldiği görülür.', style='List Bullet'
    )

    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/ters_ucgen.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"({performans_tablosu_listesi[0][0].title()} Edinme Hunisi)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        'Edinme hunisi incelendiğinde, alana gelen kişilerin', style='List Bullet'
    )
    p = document.add_paragraph(f"o 15 saniye ve üzeri alanında {edinme_hunisi_yuzde[2][1]}%.")
    p = document.add_paragraph(f"o 10 saniye ve üzeri alanında {edinme_hunisi_yuzde[1][1]}%.")
    p = document.add_paragraph(f"o 3 saniye ve üzeri alanında {edinme_hunisi_yuzde[0][1]}% zaman geçirilmektedir.")



    p = document.add_paragraph(f"{performans_tablosu_listesi[0][0].title()} alanının performans değişimi incelendiğinde;")

    document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/tablo-5.png"), width=Inches(6) , height=Inches(2))
    last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(f"({performans_tablosu_listesi[0][0].title()} Alanı Yüzdelik Performans Değişimi Grafiği)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER









    print("simdi son")

    degerler = performans_tablosu_listesi[1][3]



    ilgi_orani_artisi_2_8 = performans_tablosu_listesi[1][9]

    def ilgi_orani_durumu():
        if ilgi_orani_artisi_2_8> 0:
            return "artmıştır."
        else:
            return "azalmıştır."
    def ilgi_orani_hali():
        if ilgi_orani_artisi_2_8> 0:
            return "artarken"
        else:
            return "azalirken"


    ilgi_orani_durumu_2_8 = ilgi_orani_durumu()

    kişi_sayisi_artisi_2_8 = performans_tablosu_listesi[1][5]
    gecirilen_sure_artisi_2_8 = performans_tablosu_listesi[1][7]
    yogunluk_artisi_2_8 = performans_tablosu_listesi[1][7]

    def ilk_kisim2_8_(kişi_sayisi_artisi_2_8,gecirilen_sure_artisi_2_8,yogunluk_artisi_2_8):

        if kişi_sayisi_artisi_2_8>0:
            if gecirilen_sure_artisi_2_8 > 0:
                if yogunluk_artisi_2_8 > 0:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre hem kişi sayısının hem de geçirilen sürenin sırasıyla {kişi_sayisi_artisi_2_8}% ve {gecirilen_sure_artisi_2_8}% artmasına bağlı olarak yoğunluk {yogunluk_artisi_2_8}% artmıştır.',1]
                else:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre hem kişi sayısının hem de geçirilen sürenin sırasıyla {kişi_sayisi_artisi_2_8}% ve {gecirilen_sure_artisi_2_8}% artmasına bağlı olarak yoğunluk {yogunluk_artisi_2_8}% azalmıştır.',0]
            else:
                if yogunluk_artisi_2_8 > 0:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre kişi sayısı {kişi_sayisi_artisi_2_8}% artmıştır, geçirilen süre  {gecirilen_sure_artisi_2_8}% ile azalmıştır bunların etkisiyle; yoğunluk {yogunluk_artisi_2_8}% artmıştır.',1]
                else:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre kişi sayısı {kişi_sayisi_artisi_2_8}% artmıştır, geçirilen süre  {gecirilen_sure_artisi_2_8}% ile azalmıştır bunların etkisiyle; yoğunluk {yogunluk_artisi_2_8}% azalmıştır.',0]

        else:
            if gecirilen_sure_artisi_2_8 > 0:
                if yogunluk_artisi_2_8 > 0:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre kişi sayısı {kişi_sayisi_artisi_2_8}% ile azalmıştır, geçirilen süre  {gecirilen_sure_artisi_2_8}% ile artmıştır bunların etkisiyle; yoğunluk {yogunluk_artisi_2_8}% artmıştır.',1]
                else:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre kişi sayısı {kişi_sayisi_artisi_2_8}% ile azalmıştır, geçirilen süre  {gecirilen_sure_artisi_2_8}% ile artmıştır bunların etkisiyle; yoğunluk {yogunluk_artisi_2_8}% azalmıştır.',0]

            else:
                if yogunluk_artisi_2_8 > 0:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre kişi sayısı {kişi_sayisi_artisi_2_8}% ile azalmıştır, geçirilen süre  {gecirilen_sure_artisi_2_8}% ile azalmıştır bunların etkisiyle; yoğunluk {yogunluk_artisi_2_8}% artmıştır.',1]
                else:
                    return [f'İlgili alanda, bir önceki tarih aralığına göre kişi sayısı {kişi_sayisi_artisi_2_8}% ile azalmıştır, geçirilen süre  {gecirilen_sure_artisi_2_8}% ile azalmıştır bunların etkisiyle; yoğunluk {yogunluk_artisi_2_8}% azalmıştır.',0]

    ilk_kisim2_8_yazi = ilk_kisim2_8_(kişi_sayisi_artisi_2_8,gecirilen_sure_artisi_2_8,yogunluk_artisi_2_8)

    if ilk_kisim2_8_yazi[1]== True:
        paragraf_2 = 'Yoğunluktaki genel artışın satışı da buna yakın bir yüzdeyle arttırması beklenir.'
    else:
        paragraf_2 = 'Yoğunluktaki genel azalışın satışı buna yakın bir yüzdeyle azaltması beklenir.'

    if ilk_kisim2_8_yazi[1]== True:
        paragraf_3 = 'artarken'
    else:
        paragraf_3 = 'azalırken'

    ilgi_hali_2_8 = ilgi_orani_hali()

    def yogunluk_orani_hali():
        if ilk_kisim2_8_yazi[1] == 1:
            return "artarken"
        else:
            return "azalirken"

    yogunluk_hali_2_8 = yogunluk_orani_hali()


    document.add_paragraph(
        f'{ilk_kisim2_8_yazi[0]}', style='List Bullet'
    )
    document.add_paragraph(
        f'{paragraf_2}', style='List Bullet'
    )

    p = document.add_paragraph(f"Satışlar arttıysa;")

    p = document.add_paragraph(f"o Yoğunluk {paragraf_3} satışın da artması, alandaki potansiyelin yükselmekte olduğunu gösterir.")
    p = document.add_paragraph(f"o Satış hedefini daha da yükseltmeye yönelik aksiyon alınmalıdır.")
    p = document.add_paragraph(f"o Önceki tarih aralığına göre ilgi oranı da {ilgi_orani_artisi_2_8}% {ilgi_orani_durumu_2_8}. ")
    p = document.add_paragraph(f"o Bu alanda ilgi oranı {ilgi_hali_2_8} ve yoğunluk {yogunluk_hali_2_8} satış da arttığına göre bu alanda geçirilen sürenin satışa olan etkisi incelenmelidir.")
    p = document.add_paragraph(f"o Satışı daha da arttırmak için ziyaretçilerin alanda daha fazla zaman geçirmesi sağlanabilir.")
    p = document.add_paragraph(f"o Bunun için; personel ilgisi, ürün fiyat dengesi, ürün çeşitliliği gözden geçirilebilinir ")





    ##-sayfa-2.9_____________________________________________________

    """p = document.add_paragraph(f"")
    p.add_run(f"Kategori Karşılaştırması").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER"""

    document.add_page_break ()
    document.add_heading('2.9 Kategori Karşılaştırması', level=1)



    def populate_karsilastirma(kategori_karsilasmasi_1, kategori_karsilasmasi_2,magaza_id_no):


        p = document.add_paragraph(f"")
        p.add_run(f"{kategori_karsilasmasi_1} & {kategori_karsilasmasi_2}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kategori_1 = request1.main_2_7_performans_tablosu_double_isim(magaza_id_no, ilk_tarih , son_tarih,kategori_karsilasmasi_1)
        kategori_2 = request1.main_2_7_performans_tablosu_double_isim ( magaza_id_no, ilk_tarih, son_tarih,
                                                                        kategori_karsilasmasi_2 )

        kategori_1_kisi_sayisi = kategori_1[1][3]  # kisi sayisi
        kategori_1_sure = kategori_1[1][6]  # sure
        kategori_1_m2 = kategori_1[1][10]  # sure
        kategori_1_m2_orani = kategori_1[1][11]  # sure
        kategori_1_yogunluk = kategori_1[1][0]
        kategori_1_m2_bolu_yogunluk = float(kategori_1_m2)/float(kategori_1_yogunluk)


        kategori_2_kisi_sayisi = kategori_2[1][3]  # kisi sayisi
        kategori_2_sure = kategori_2[1][6]  # sure
        kategori_2_m2 = kategori_2[1][10]  # sure
        kategori_2_m2_orani = kategori_2[1][11]  # sure
        kategori_2_yogunluk = kategori_2[1][0]
        kategori_2_m2_bolu_yogunluk = float ( kategori_2_m2 ) / float ( kategori_2_yogunluk )

        print("su sayilara bakk")
        print(kategori_1_m2_bolu_yogunluk)
        print ( kategori_2_m2_bolu_yogunluk )


        def yazilar(kategori_karsilasmasi_1,kategori_1_kisi_sayisi,kategori_1_sure,kategori_karsilasmasi_2,kategori_2_kisi_sayisi,kategori_2_sure):
            if kategori_1_kisi_sayisi> kategori_2_kisi_sayisi:
                print("süre ağırlıklı")
                if kategori_1_sure > kategori_2_sure:
                    return f'{kategori_karsilasmasi_1.title()} alanında yoğunluğu oluşturan asıl parametre kişi sayısı ve ortalama geçirilen süredir. {kategori_karsilasmasi_2} alanına daha çok odaklanılmalıdır. Bu alanda geçirilen ortalama süre sadece {kategori_1_sure} saniye, ortalama kişi sayısı {kategori_1_kisi_sayisi} kişidir.'
                else:
                    return f'{kategori_karsilasmasi_1.title()} alanında yoğunluğu oluşturan asıl parametre kişi sayısıdır. Bu alanda yoğunluğu arttırmak için süreyi arttırmaya odaklanmak gerekir. Bu alanda geçirilen ortalama süre sadece {kategori_1_sure} saniyedir.'

            elif kategori_1_sure > kategori_2_sure:
                return f'{kategori_karsilasmasi_1.title()} alanında yoğunluğu oluşturan asıl parametre süre ağırlıklıdır. Bu alanda yoğunluğu arttırmak için kişi sayısını arttırmaya odaklanmak gerekir. Bu alandaki ortalama kişi sayısı {kategori_1_sure} saniyedir.'
            else:
                return f'{kategori_karsilasmasi_2.title()} alanına göre kişi veya süre ağırlıklı değildir. Bu alanda yoğunluğu arttırmak için kişi sayısını ve süreyi arttırmaya odaklanmak gerekir.'

        ilk_part = yazilar(kategori_karsilasmasi_1,kategori_1_kisi_sayisi,kategori_1_sure,kategori_karsilasmasi_2,kategori_2_kisi_sayisi,kategori_2_sure)
        ikici_part = yazilar(kategori_karsilasmasi_2,kategori_2_kisi_sayisi,kategori_2_sure,kategori_karsilasmasi_1,kategori_1_kisi_sayisi,kategori_1_sure)


        def m2_karsilastirma(kategori_karsilasmasi_1,kategori_1_m2,kategori_karsilasmasi_2,kategori_2_m2):
            if kategori_1_m2 >kategori_2_m2:
                return(f'Mağaza içerisinde {kategori_karsilasmasi_1.title()} alanının m2 orani büyüktür, fark {("{:.2f}".format(round(kategori_1_m2-kategori_2_m2, 2)))} kadardır.') #{("{:.2f}".format(round(kategori_1_m2-kategori_2_m2, 2)))}
            elif kategori_2_m2 >kategori_1_m2:
                return ( f'Mağaza içerisinde {kategori_karsilasmasi_2.title()} alanının m2 orani küçüktür,fark {("{:.2f}".format(round(kategori_2_m2-kategori_1_m2, 2)))} kadardır.' )
            elif kategori_1_m2 +2 >kategori_2_m2 or kategori_1_m2 -2 >kategori_2_m2:
                return ( f"Mağaza içerisinde her iki alana ayrılan metre kareler yaklaşık olarak eşittir." )

        m2_karsilartirma_1 = m2_karsilastirma(kategori_karsilasmasi_1,kategori_1_m2,kategori_karsilasmasi_2,kategori_2_m2)

        def yogunluk_karsilastirma(kategori_karsilasmasi_1,kategori_1_yogunluk,kategori_karsilasmasi_2,kategori_2_yogunluk):
            if kategori_1_yogunluk >kategori_2_yogunluk:
                return(f'{kategori_karsilasmasi_1.title()} alanı, {kategori_karsilasmasi_2.title()} alanının yaklaşık olarak {("{:.2f}".format(round(kategori_1_yogunluk/kategori_2_yogunluk, 2)))} katı yoğunluğa sahiptir.')
            else:
                return (f'{kategori_karsilasmasi_2.title()} alanı, {kategori_karsilasmasi_1.title()} alanının yaklaşık olarak {("{:.2f}".format(round(kategori_2_yogunluk/kategori_1_yogunluk, 2)))} katı yoğunluğa sahiptir.')

        yogunluk_karsilastirma_1 = yogunluk_karsilastirma(kategori_karsilasmasi_1,kategori_1_yogunluk,kategori_karsilasmasi_2,kategori_2_yogunluk)



        document.add_picture(os.path.join(BASE_DIR, f"{magza_statik_dosya_location}/tablo-4.png"), width=Inches(6) , height=Inches(0.75))
        last_paragraph = document.paragraphs[-1]#resimleri ortalamak icin
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph(f"({kategori_karsilasmasi_1.title()} ve {kategori_karsilasmasi_2.title()} Alanlarının Yoğunlukları)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        p = document.add_paragraph(f"")
        p.add_run(f'Mağaza içerisindeki {kategori_karsilasmasi_1.title()} ve {kategori_karsilasmasi_1.title()} alanlarının yoğunlukları, metre kareleri ve metre kare başına düşen yoğunluk yüzdeleri kıyaslanmıştır.').italic = True


        document.add_paragraph(
            f'{m2_karsilartirma_1}', style='List Bullet'
        )
        document.add_paragraph(
            f'{yogunluk_karsilastirma_1}', style='List Bullet'
        )
        document.add_paragraph(
            'Yoğunluk paylarının, satış paylarına paralel olması beklenir.', style='List Bullet' # burayi gelistir
        )
        document.add_paragraph(
            f'Metre kare başına düşen yoğunluklar incelendiğinde, {kategori_karsilasmasi_2.title()} alanının metre kare başına düşen yoğunluğu, {kategori_karsilasmasi_1.title()}’dan daha büyüktür ve bu oran 1’in üzerindedir.', style='List Bullet'
        )
        document.add_paragraph(
            f'Yani 1 birim metre kareye 1.15 birim yoğunluk düşmektedir. Alan yoğunluk bakımından verimli kullanılmıştır.', style='List Bullet'
        )
        document.add_paragraph(
            f'{kategori_karsilasmasi_2.title()} alanında yoğunluk/metre kare oranı 1’den küçüktür. Buna göre bu alan, yoğunluk bazında verimli kullanılamamıştır. Alana ayrılan metre karenin büyük geldiği düşünülüyorsa alan kademeli olarak küçültülmeli, metre kare küçülürken satış azalmıyorsa alan küçültülmelidir.', style='List Bullet'
        )
        document.add_paragraph(
            f'Her iki alanda da yoğunluğu ve dolayısıyla satışları arttırmak için alana gelen kişi sayısını ve geçirilen süreyi arttırmak için aksiyon alınmalıdır.', style='List Bullet'
        )
        document.add_paragraph(
            f'Yoğunluktaki artışın satışı da arttırması beklenir.', style='List Bullet'
        )

        document.add_picture ( os.path.join ( BASE_DIR, f"{magza_statik_dosya_location}/tablo-2.png" ),
                               width=Inches ( 6 ), height=Inches ( 0.75 ) )
        last_paragraph = document.paragraphs[-1]  # resimleri ortalamak icin
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_picture ( os.path.join ( BASE_DIR, f"{magza_statik_dosya_location}/tablo-1.png" ),
                               width=Inches ( 6 ), height=Inches ( 0.75 ) )
        last_paragraph = document.paragraphs[-1]  # resimleri ortalamak icin
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph ( f"({kategori_karsilasmasi_1.title()} ve {kategori_karsilasmasi_2.title()} Tablosu)" )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER





        p = document.add_paragraph(f"Detaya inildiğinde;")
        document.add_paragraph(
            f'{ilk_part}', style='List Bullet'
        )

        document.add_paragraph(
            f'{ikici_part}', style='List Bullet'
        )
        #document.add_page_break ()


    for karsilastirilacak in karsilastirilacak_isim_listesi:
        populate_karsilastirma(karsilastirilacak[0],karsilastirilacak[1],magaza_id_no)


    ###-11_______________________________________________


    p = document.add_paragraph(f"Rapor ile ilgili sorularınızı ve yorumlarınızı hello@udentify.co adresine iletebilirsiniz. İyi çalışmalar!")





    #document.add_picture('monty-truth.png', width=Inches(1.25))

    #sayfa-3.1
    document.add_page_break()
    document.add_heading('3.0 Sonuç', 0)

    document.add_heading('3.1 Analizin Sonucu', level=1)


    p = document.add_paragraph(f"Bu raporda {magaza_adi} mağazası detaylı olarak incelenmiştir. Alınacak muhtemel aksiyonlar da gösterilmiştir.\nBu aksiyonların uygulanıp uygulanmadığını Udentify tarafına bildirdiğinizde detaylı olarak analizlerle değişimlerin başarısını ölçümleyip size raporlamak isteriz.")




    print("________Bitti")

    document.add_page_break()

    document.save(f"{magaza_dosyasi_ismi}")



#refresh() # burayi degistirdin



#start_writing_on_docx(magza_statik_dosya_location,magaza_dosyasi_ismi,200,ilk_tarih,son_tarih,gereken_dosya_ismi)


#start_writing_on_docx("Under Armour","under atmour akasyya")


#burasi test kismi

magaza_adi_listesi = [["Under Armour","Akasya",240,"Under Armour Akasya",240],["Under Armour","Zorlu Center",239,"Under Armour Zorlu Center",240],["Under Armour","İstinye Park",228,"Under Armour Istinye Park",240]]
#
# start_writing_on_docx(firma,magza_statik_dosya_location_ismi,240,ilk_tarih,son_tarih,161,240,[["a","b"],["c","d"]])  ##birincisi firma adi ile dosya konumunu bulmaya yariyor   ,ikincisi statik dosya locationlarini bulmaya yariyor
#
global_test =  True
#
# ilk_tarih = "30/10/2020"
# son_tarih = "13/11/2020"
#

# yeni_liste =[["Under Armour","Under Armour Zorlu Center",239,ilk_tarih,son_tarih,160,[["BOYS","GIRLS"],["WOMEN'S RUN","MEN'S RUN"]]],["Under Armour","Under Armour Akasya",240,ilk_tarih,son_tarih,161,[["BOYS","GIRLS"],["WOMEN'S RUN","MEN'S RUN"]]],["Under Armour","Under Armour Istinye Park",228,ilk_tarih,son_tarih,149,[["BOYS","GIRLS"],["WOMEN'S RUN","MEN'S RUN"]]]]
#
# for liste in yeni_liste:
#     start_writing_on_docx ( liste[0], liste[1], liste[2], liste[3], liste[4], liste[5],
#                             liste[6] )


#start_writing_on_docx("Under Armour","Under Armour Zorlu Center",239,ilk_tarih,son_tarih,160,[["BOYS","GIRLS"],["WOMEN'S RUN","MEN'S RUN"]])
#start_writing_on_docx("Under Armour","Under Armour Akasya",240,ilk_tarih,son_tarih,161,[["BOYS","GIRLS"],["WOMEN'S RUN","MEN'S RUN"]])
#start_writing_on_docx("Under Armour","Under Armour Istinye Park",228,ilk_tarih,son_tarih,149,[["BOYS","GIRLS"],["WOMEN'S RUN","MEN'S RUN"]])





#start_writing_on_docx("Suwen","Suwen Viaport",307,ilk_tarih,son_tarih,189,[["TAYT","ÇORAP"],["ATLET","ERKEK REYONU"]]) #density kisimlari error veriyo cikarmadinn daha



